#!/usr/bin/env python3
"""Render a polished Korean legal-opinion DOCX from research markdown.

Differs from `render-docx.py` (basic MVP renderer) by adding:

- Cover page with title, recipient, date, classification banner.
- Page numbers in footer (PAGE / NUMPAGES field).
- Section break before each top-level §X heading.
- Heading numbering with Korean legal convention (제X장, X.Y., X.Y.Z.).
- Tables with shaded header row, full borders, alternating-row treatment.
- Justified body paragraphs with proper Korean line spacing.
- Footer with confidentiality classification.
- Endnote-style 각주 (footnotes) section at document end.
- Pinpoint citation styling — verbatim 법령 quotations rendered as indented italic.
- Markdown footnote markers (`[^xxx]`) converted to superscript numerals
  with a master 각주 list at the end.

Single-purpose script for polished Korean legal-opinion-style reports;
generalisable to any well-structured Korean legal-opinion-style markdown.

Usage:

    python3 scripts/render-legal-opinion-docx.py \
        path/to/report.md \
        path/to/report.docx \
        --title "규제 검토 보고서" \
        --recipient "사내 법무팀 귀중" \
        --date "2026년 5월 7일" \
        --classification "CONFIDENTIAL — INTERNAL LEGAL REVIEW" \
        --author "Legal Research Agent"
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.shared import Inches, Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsmap
except ImportError as exc:
    print(f"FAIL: python-docx is required: {exc}", file=sys.stderr)
    raise


# ---------------------------------------------------------------------------
# Page / typography defaults — Korean legal-opinion convention
# ---------------------------------------------------------------------------

PAGE = {
    "width": Mm(210),
    "height": Mm(297),
    "top": Mm(28),
    "bottom": Mm(25),
    "left": Mm(28),
    "right": Mm(25),
}

FONT_BODY_HANGUL = "맑은 고딕"
FONT_BODY_LATIN = "Times New Roman"
FONT_HEADING_HANGUL = "맑은 고딕"
FONT_HEADING_LATIN = "Times New Roman"

SIZE_COVER_TITLE = Pt(22)
SIZE_COVER_SUBTITLE = Pt(13)
SIZE_COVER_META = Pt(11)
SIZE_H1 = Pt(15)
SIZE_H2 = Pt(13)
SIZE_H3 = Pt(11.5)
SIZE_H4 = Pt(11)
SIZE_BODY = Pt(10.5)
SIZE_QUOTE = Pt(10)
SIZE_TABLE = Pt(9.5)
SIZE_TABLE_HEADER = Pt(9.5)
SIZE_FOOTER = Pt(9)
SIZE_FOOTNOTE = Pt(9)

LINE_SPACING_BODY = 1.45
LINE_SPACING_TIGHT = 1.2

COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_INK = RGBColor(0x12, 0x18, 0x2B)
COLOR_GRAY_DARK = RGBColor(0x4A, 0x55, 0x68)
COLOR_GRAY_MID = RGBColor(0x94, 0xA3, 0xB8)
COLOR_GRAY_LIGHT = RGBColor(0xE2, 0xE8, 0xF0)
COLOR_TABLE_HEADER_BG = "1F2937"  # slate-800 hex (no #)
COLOR_TABLE_ALT_BG = "F8FAFC"     # slate-50
COLOR_BANNER_BG = "FEF3C7"
COLOR_BANNER_BORDER = "F59E0B"


# ---------------------------------------------------------------------------
# Markdown parsing — purpose-built for the report's structure
# ---------------------------------------------------------------------------

INLINE_FOOTNOTE_RE = re.compile(r"\[\^([^\]]+)\]")
FOOTNOTE_DEF_RE = re.compile(r"^\[\^([^\]]+)\]:\s*(.*)$")
INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITALIC_RE = re.compile(r"(?<![\*])\*(?!\s)([^*]+?)\*(?![\*])")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")
INLINE_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


@dataclass
class Block:
    kind: str
    level: int = 0
    text: str = ""
    items: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


def parse_markdown(text: str) -> tuple[list[Block], dict[str, str]]:
    """Parse markdown into a flat list of Blocks plus a footnote dictionary."""
    lines = text.splitlines()
    blocks: list[Block] = []
    footnotes: dict[str, str] = {}
    i = 0

    def is_table_row(s: str) -> bool:
        return s.lstrip().startswith("|") and s.rstrip().endswith("|")

    def is_table_separator(s: str) -> bool:
        return bool(re.match(r"^\s*\|[\s:|-]+\|\s*$", s))

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Footnote definition
        m = FOOTNOTE_DEF_RE.match(line)
        if m:
            key, body = m.group(1), m.group(2)
            # multi-line footnote: continue while next lines are indented or empty
            j = i + 1
            while j < len(lines) and (lines[j].startswith("    ") or lines[j].startswith("\t")):
                body += " " + lines[j].strip()
                j += 1
            footnotes[key] = body.strip()
            i = j
            continue

        # Heading
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text_ = line[level:].strip()
            blocks.append(Block(kind="heading", level=level, text=text_))
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            blocks.append(Block(kind="hr"))
            i += 1
            continue

        # Table (header + separator + rows)
        if is_table_row(line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = []
            # header
            rows.append(_split_table_row(line))
            i += 2  # skip separator
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(_split_table_row(lines[i]))
                i += 1
            blocks.append(Block(kind="table", rows=rows))
            continue

        # Block quote
        if line.startswith(">"):
            quoted = []
            while i < len(lines) and lines[i].startswith(">"):
                quoted.append(lines[i].lstrip("> ").rstrip())
                i += 1
            blocks.append(Block(kind="quote", text="\n".join(quoted).strip()))
            continue

        # Code fence
        if line.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # closing fence
            blocks.append(Block(kind="code", text="\n".join(buf)))
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]).rstrip())
                i += 1
            blocks.append(Block(kind="ul", items=items))
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]).rstrip())
                i += 1
            blocks.append(Block(kind="ol", items=items))
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Paragraph (collect until blank line / structural element)
        para = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if nxt.startswith(("#", "|", ">", "```")):
                break
            if re.match(r"^\s*([-*]|\d+\.)\s+", nxt):
                break
            if FOOTNOTE_DEF_RE.match(nxt):
                break
            if nxt.strip() in ("---", "***", "___"):
                break
            para.append(nxt)
            i += 1
        blocks.append(Block(kind="paragraph", text=" ".join(para).strip()))

    return blocks, footnotes


def _split_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


# ---------------------------------------------------------------------------
# Inline span rendering helpers
# ---------------------------------------------------------------------------


def _add_runs_with_inlines(paragraph, text: str, footnote_index: dict[str, int],
                           base_size: Pt = SIZE_BODY, italic: bool = False, bold: bool = False):
    """Render inline markdown into runs (bold, italic, code, links, footnotes)."""

    # First, capture footnote markers and replace with sentinels so they don't
    # get consumed by other inline patterns.
    sentinels: dict[str, str] = {}

    def _make_sentinel(idx: int) -> str:
        return f"FN{idx}"

    fn_iter = list(INLINE_FOOTNOTE_RE.finditer(text))
    for m in fn_iter:
        key = m.group(1)
        if key in footnote_index:
            idx = footnote_index[key]
        else:
            idx = len(footnote_index) + 1
            footnote_index[key] = idx
        sentinels[m.group(0)] = _make_sentinel(idx)
    for original, sent in sentinels.items():
        text = text.replace(original, sent)

    # Walk through text emitting runs
    pos = 0
    pattern = re.compile(
        r"(FN\d+)|"
        r"(\*\*[^*]+\*\*)|"
        r"(`[^`]+`)|"
        r"(\[[^\]]+\]\([^)]+\))|"
        r"((?<![\*])\*(?!\s)[^*]+\*(?![\*]))"
    )
    while pos < len(text):
        m = pattern.search(text, pos)
        if not m:
            _emit_run(paragraph, text[pos:], base_size, italic=italic, bold=bold)
            break
        if m.start() > pos:
            _emit_run(paragraph, text[pos:m.start()], base_size, italic=italic, bold=bold)
        token = m.group(0)
        if token.startswith("FN"):
            num = int(token[3:-1])
            run = paragraph.add_run(str(num))
            run.font.superscript = True
            _set_run_fonts(run, base_size, bold=bold)
        elif token.startswith("**") and token.endswith("**"):
            inner = token[2:-2]
            _emit_run(paragraph, inner, base_size, italic=italic, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            inner = token[1:-1]
            run = paragraph.add_run(inner)
            run.font.name = "Consolas"
            r = run._element.rPr
            if r is None:
                r = OxmlElement("w:rPr")
                run._element.insert(0, r)
            rfonts = r.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                r.append(rfonts)
            rfonts.set(qn("w:ascii"), "Consolas")
            rfonts.set(qn("w:hAnsi"), "Consolas")
            rfonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)
        elif token.startswith("[") and "](" in token:
            link_m = INLINE_LINK_RE.match(token)
            if link_m:
                label, _ = link_m.group(1), link_m.group(2)
                _emit_run(paragraph, label, base_size, italic=italic, bold=bold)
        else:  # italic
            inner = token[1:-1]
            _emit_run(paragraph, inner, base_size, italic=True, bold=bold)
        pos = m.end()


def _emit_run(paragraph, text: str, size: Pt, *, italic: bool = False, bold: bool = False):
    if not text:
        return
    run = paragraph.add_run(text)
    run.italic = italic
    run.bold = bold
    _set_run_fonts(run, size, bold=bold, italic=italic)


def _set_run_fonts(run, size: Pt, *, bold: bool = False, italic: bool = False):
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rfonts = rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_BODY_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_BODY_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_BODY_HANGUL)
    rfonts.set(qn("w:cs"), FONT_BODY_LATIN)


# ---------------------------------------------------------------------------
# Document building
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # remove any existing shading element
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "94A3B8", size_eighths: int = 4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        old = tcBorders.find(qn(f"w:{edge}"))
        if old is not None:
            tcBorders.remove(old)
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_eighths))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)


def _add_section_break(doc, *, restart_page: bool = False):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _add_page_field(paragraph, *, total: bool = False):
    instr_text = "NUMPAGES" if total else "PAGE"
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {instr_text} "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    rPr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), FONT_BODY_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_BODY_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_BODY_HANGUL)
    rPr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def _set_page_margins(section):
    section.page_width = PAGE["width"]
    section.page_height = PAGE["height"]
    section.top_margin = PAGE["top"]
    section.bottom_margin = PAGE["bottom"]
    section.left_margin = PAGE["left"]
    section.right_margin = PAGE["right"]


def _setup_section_with_footer(section, *, classification: str):
    _set_page_margins(section)
    section.different_first_page_header_footer = False

    # Footer content: classification (left) + page X / Y (right)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(section.page_width - section.left_margin - section.right_margin)
    run = p.add_run(classification)
    _set_run_fonts(run, SIZE_FOOTER)
    run.font.color.rgb = COLOR_GRAY_DARK
    p.add_run("\t")
    _add_page_field(p)
    page_run_separator = p.add_run(" / ")
    _set_run_fonts(page_run_separator, SIZE_FOOTER)
    page_run_separator.font.color.rgb = COLOR_GRAY_DARK
    _add_page_field(p, total=True)

    # Apply same font to all footer runs
    for r in p.runs:
        if r.font.size is None:
            _set_run_fonts(r, SIZE_FOOTER)
            r.font.color.rgb = COLOR_GRAY_DARK


def _add_cover_page(doc, *, title: str, subtitle: str, recipient: str,
                    date: str, classification: str, author: str):
    # Top spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Classification banner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(classification)
    run.bold = True
    _set_run_fonts(run, Pt(10))
    run.font.color.rgb = RGBColor(0x78, 0x35, 0x0F)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), COLOR_BANNER_BG)
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    _set_run_fonts(run, SIZE_COVER_TITLE)
    run.font.color.rgb = COLOR_INK
    p.paragraph_format.space_before = Pt(60)

    # Subtitle (English)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.italic = True
        _set_run_fonts(run, SIZE_COVER_SUBTITLE)
        run.font.color.rgb = COLOR_GRAY_DARK
        p.paragraph_format.space_after = Pt(24)

    # Horizontal rule (visual)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 30)
    _set_run_fonts(run, Pt(10))
    run.font.color.rgb = COLOR_GRAY_MID

    # Meta block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run(f"수신: {recipient}")
    _set_run_fonts(run, SIZE_COVER_META)
    run.bold = True
    run.font.color.rgb = COLOR_INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"작성일: {date}")
    _set_run_fonts(run, SIZE_COVER_META)
    run.font.color.rgb = COLOR_INK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"작성: {author}")
    _set_run_fonts(run, SIZE_COVER_META)
    run.font.color.rgb = COLOR_INK

    # Disclaimer block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run(
        "본 보고서는 내부 법률 검토 지원용 보고서로,\n"
        "특정 거래·소송에 대한 법률자문을 대체하지 아니합니다.\n"
        "중요한 의사결정 전 담당 전문가의 추가 검토를 권장합니다."
    )
    run.italic = True
    _set_run_fonts(run, Pt(9.5))
    run.font.color.rgb = COLOR_GRAY_DARK

    _add_section_break(doc)


# ---------------------------------------------------------------------------
# Body element rendering
# ---------------------------------------------------------------------------


@dataclass
class HeadingNumberer:
    h1: int = 0
    h2: int = 0
    h3: int = 0
    h4: int = 0

    def number_for(self, level: int) -> str:
        if level == 1:
            return ""
        if level == 2:
            self.h2 += 1
            self.h3 = 0
            self.h4 = 0
            return f"{self.h2}."
        if level == 3:
            self.h3 += 1
            self.h4 = 0
            return f"{self.h2}.{self.h3}."
        if level == 4:
            self.h4 += 1
            return f"{self.h2}.{self.h3}.{self.h4}."
        return ""


def _render_heading(doc, level: int, text: str, numberer: HeadingNumberer,
                    footnote_index: dict[str, int]):
    if level == 1:
        # Top-level document title — already covered by cover page; skip
        return
    if level == 2:
        # Top-level section (§X) — page break before, prominent
        _add_section_break(doc)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        # Underline-style bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "0F172A")
        pBdr.append(bottom)
        pPr.append(pBdr)
        run = p.add_run(text.lstrip("0123456789. "))
        run.bold = True
        _set_run_fonts(run, SIZE_H1)
        run.font.color.rgb = COLOR_INK
        return

    label = numberer.number_for(level)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12) if level == 2 else Pt(8)
    p.paragraph_format.space_after = Pt(4)
    if label:
        prefix_run = p.add_run(label + " ")
        prefix_run.bold = True
        _set_run_fonts(prefix_run, SIZE_H2 if level <= 3 else SIZE_H4)
        prefix_run.font.color.rgb = COLOR_INK
    body_run = p.add_run(text.lstrip("0123456789. "))
    body_run.bold = True
    if level == 3:
        _set_run_fonts(body_run, SIZE_H2)
    elif level == 4:
        _set_run_fonts(body_run, SIZE_H3)
    else:
        _set_run_fonts(body_run, SIZE_H4)
    body_run.font.color.rgb = COLOR_INK


def _render_paragraph(doc, text: str, footnote_index: dict[str, int]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING_BODY
    p.paragraph_format.space_after = Pt(6)
    _add_runs_with_inlines(p, text, footnote_index, base_size=SIZE_BODY)


def _render_quote(doc, text: str, footnote_index: dict[str, int]):
    for line in text.split("\n"):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(8)
        p.paragraph_format.right_indent = Mm(8)
        p.paragraph_format.space_after = Pt(2)
        # left border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "94A3B8")
        pBdr.append(left)
        pPr.append(pBdr)
        # subtle gray background
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F8FAFC")
        pPr.append(shd)
        _add_runs_with_inlines(p, line.strip(), footnote_index, base_size=SIZE_QUOTE, italic=True)


def _render_list(doc, items: list[str], *, ordered: bool, footnote_index: dict[str, int]):
    style_name = "List Number" if ordered else "List Bullet"
    for item in items:
        try:
            p = doc.add_paragraph(style=style_name)
        except KeyError:
            # fall back if style not found
            p = doc.add_paragraph()
            marker = "• " if not ordered else ""
            p.add_run(marker)
        p.paragraph_format.line_spacing = LINE_SPACING_BODY
        p.paragraph_format.space_after = Pt(2)
        _add_runs_with_inlines(p, item, footnote_index, base_size=SIZE_BODY)


def _render_table(doc, rows: list[list[str]], footnote_index: dict[str, int]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # iterate
    for r_idx, row_cells in enumerate(rows):
        # pad short rows
        padded = row_cells + [""] * (cols - len(row_cells))
        for c_idx, cell_text in enumerate(padded):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # remove default paragraph
            for p in list(cell.paragraphs):
                p.clear()
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            is_header = r_idx == 0
            text_clean = cell_text.replace("<br>", "\n")
            for sub_idx, sub_line in enumerate(text_clean.split("\n")):
                if sub_idx > 0:
                    p = cell.add_paragraph()
                if is_header:
                    run = p.add_run(sub_line.strip())
                    run.bold = True
                    _set_run_fonts(run, SIZE_TABLE_HEADER)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    _add_runs_with_inlines(p, sub_line.strip(), footnote_index,
                                            base_size=SIZE_TABLE)
            # shading
            if is_header:
                _set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
            elif r_idx % 2 == 0:
                _set_cell_shading(cell, COLOR_TABLE_ALT_BG)
            _set_cell_borders(cell)


def _render_code(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rfonts.set(qn("w:eastAsia"), "Consolas")
    rPr.append(rfonts)


def _render_hr(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "94A3B8")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Footnote (각주) rendering at document end
# ---------------------------------------------------------------------------


def _render_footnotes(doc, footnotes: dict[str, str], footnote_index: dict[str, int]):
    if not footnote_index:
        return
    _add_section_break(doc)

    # Section heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F172A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run("각주 (Endnotes)")
    run.bold = True
    _set_run_fonts(run, SIZE_H1)
    run.font.color.rgb = COLOR_INK

    # Sort by index
    ordered_keys = sorted(footnote_index.keys(), key=lambda k: footnote_index[k])
    for key in ordered_keys:
        idx = footnote_index[key]
        body = footnotes.get(key, "[footnote text not found]")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(7)
        p.paragraph_format.first_line_indent = Mm(-7)
        p.paragraph_format.space_after = Pt(2)
        num_run = p.add_run(f"{idx}. ")
        num_run.bold = True
        _set_run_fonts(num_run, SIZE_FOOTNOTE)
        num_run.font.color.rgb = COLOR_GRAY_DARK
        # Body — strip inline markdown footnote references that point to other notes
        # to avoid recursive numbering; use base_size SIZE_FOOTNOTE
        body_clean = INLINE_FOOTNOTE_RE.sub("", body)
        _add_runs_with_inlines(p, body_clean, {}, base_size=SIZE_FOOTNOTE)


# ---------------------------------------------------------------------------
# Top-level orchestration
# ---------------------------------------------------------------------------


def build_document(
    *,
    markdown_text: str,
    title: str,
    subtitle: str,
    recipient: str,
    date: str,
    classification: str,
    author: str,
) -> Document:
    blocks, footnotes = parse_markdown(markdown_text)
    doc = Document()
    section = doc.sections[0]
    _setup_section_with_footer(section, classification=classification)

    # Set base style font
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY_LATIN
    style.font.size = SIZE_BODY
    rPr = style.element.get_or_add_rPr()
    rfonts = rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_BODY_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_BODY_LATIN)
    rfonts.set(qn("w:eastAsia"), FONT_BODY_HANGUL)

    # Cover page
    _add_cover_page(
        doc,
        title=title,
        subtitle=subtitle,
        recipient=recipient,
        date=date,
        classification=classification,
        author=author,
    )

    numberer = HeadingNumberer()
    footnote_index: dict[str, int] = {}

    # Skip leading H1 (document title) and any leading metadata paragraphs that
    # duplicate cover info before the first H2.
    skip_until_h2 = True

    for block in blocks:
        if skip_until_h2:
            if block.kind == "heading" and block.level == 2:
                skip_until_h2 = False
            else:
                continue
        if block.kind == "heading":
            _render_heading(doc, block.level, block.text, numberer, footnote_index)
        elif block.kind == "paragraph":
            _render_paragraph(doc, block.text, footnote_index)
        elif block.kind == "quote":
            _render_quote(doc, block.text, footnote_index)
        elif block.kind == "ul":
            _render_list(doc, block.items, ordered=False, footnote_index=footnote_index)
        elif block.kind == "ol":
            _render_list(doc, block.items, ordered=True, footnote_index=footnote_index)
        elif block.kind == "table":
            _render_table(doc, block.rows, footnote_index)
        elif block.kind == "code":
            _render_code(doc, block.text)
        elif block.kind == "hr":
            _render_hr(doc)

    _render_footnotes(doc, footnotes, footnote_index)

    return doc


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--title", required=True)
    parser.add_argument("--subtitle", default="")
    parser.add_argument("--recipient", default="사내 법무팀 귀중")
    parser.add_argument("--date", required=True)
    parser.add_argument("--classification",
                        default="CONFIDENTIAL — INTERNAL LEGAL REVIEW ONLY")
    parser.add_argument("--author", default="Legal Research Agent")
    args = parser.parse_args(argv)

    if not args.input_md.exists():
        print(f"FAIL: input not found: {args.input_md}", file=sys.stderr)
        return 1

    text = args.input_md.read_text(encoding="utf-8")
    doc = build_document(
        markdown_text=text,
        title=args.title,
        subtitle=args.subtitle,
        recipient=args.recipient,
        date=args.date,
        classification=args.classification,
        author=args.author,
    )
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output_docx))
    print(f"OK: rendered legal-opinion DOCX → {args.output_docx} ({args.output_docx.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
