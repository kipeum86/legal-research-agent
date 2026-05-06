#!/usr/bin/env python3
"""Render standalone Markdown deliverables into a basic legal DOCX file."""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt
except ImportError as exc:  # pragma: no cover - exercised only when dependency is absent.
    print(f"FAIL: python-docx is required: {exc}", file=sys.stderr)
    raise


PAGE_CONFIGS: dict[str, dict[str, Any]] = {
    "ko": {
        "page_width": Mm(210),
        "page_height": Mm(297),
        "margins": {"top": Mm(25.4), "bottom": Mm(25.4), "left": Mm(25.4), "right": Mm(25.4)},
        "body_font_ascii": "Times New Roman",
        "body_font_east_asia": "맑은 고딕",
        "heading_font_ascii": "Times New Roman",
        "heading_font_east_asia": "맑은 고딕",
        "body_size": 11,
        "line_spacing": 1.15,
    },
    "en-us": {
        "page_width": Inches(8.5),
        "page_height": Inches(11),
        "margins": {"top": Inches(1), "bottom": Inches(1), "left": Inches(1), "right": Inches(1)},
        "body_font_ascii": "Times New Roman",
        "body_font_east_asia": "Times New Roman",
        "heading_font_ascii": "Times New Roman",
        "heading_font_east_asia": "Times New Roman",
        "body_size": 12,
        "line_spacing": 1.5,
    },
    "en-intl": {
        "page_width": Mm(210),
        "page_height": Mm(297),
        "margins": {"top": Mm(25.4), "bottom": Mm(25.4), "left": Mm(25.4), "right": Mm(25.4)},
        "body_font_ascii": "Times New Roman",
        "body_font_east_asia": "Times New Roman",
        "heading_font_ascii": "Arial",
        "heading_font_east_asia": "Arial",
        "body_size": 12,
        "line_spacing": 1.5,
    },
}

INLINE_TOKEN_RE = re.compile(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")


@dataclass
class MarkdownBlock:
    kind: str
    text: str = ""
    level: int | None = None
    rows: list[list[str]] = field(default_factory=list)
    items: list[str] = field(default_factory=list)


class RenderError(ValueError):
    """Raised when a Markdown-to-DOCX render request is invalid."""


def config_key(language: str, jurisdiction: str) -> str:
    if language == "ko":
        return "ko"
    if jurisdiction == "us":
        return "en-us"
    return "en-intl"


def render_docx(
    input_md: Path,
    output_docx: Path,
    *,
    language: str,
    jurisdiction: str = "korea",
    classification: str | None = None,
    overwrite: bool = False,
    report_path: Path | None = None,
) -> dict[str, Any]:
    validate_paths(input_md, output_docx, overwrite=overwrite)
    content = input_md.read_text(encoding="utf-8")
    blocks = parse_markdown(content)
    if not blocks:
        raise RenderError("Input Markdown did not produce any renderable blocks.")

    config = PAGE_CONFIGS[config_key(language, jurisdiction)]
    document = Document()
    configure_document(document, config, classification=classification)
    for block in blocks:
        render_block(document, block, config)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))

    report = {
        "renderer": "legal-research-agent.render-docx.v1",
        "input": str(input_md),
        "output": str(output_docx),
        "language": language,
        "jurisdiction": jurisdiction,
        "classification": classification,
        "blocks": len(blocks),
        "tables": sum(1 for block in blocks if block.kind == "table"),
        "output_bytes": output_docx.stat().st_size,
        "limitations": [
            "Native Word footnotes are not generated; footnote syntax is rendered as visible text.",
            "Native tracked changes, comments, page numbers beyond the footer field, and complex nested Markdown are not generated.",
            "Citation audit must be run separately on the Markdown source or extracted DOCX text.",
        ],
    }
    if report_path is not None:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")
    return report


def validate_paths(input_md: Path, output_docx: Path, *, overwrite: bool) -> None:
    if not input_md.exists() or not input_md.is_file():
        raise RenderError(f"Input Markdown does not exist: {input_md}")
    if input_md.suffix.lower() not in {".md", ".markdown"}:
        raise RenderError("Input must be a Markdown file with .md or .markdown extension.")
    if output_docx.suffix.lower() != ".docx":
        raise RenderError("Output must have .docx extension.")
    if output_docx.exists() and not overwrite:
        raise RenderError(f"Output already exists: {output_docx}. Use --overwrite to replace it.")


def configure_document(document: Document, config: dict[str, Any], *, classification: str | None) -> None:
    section = document.sections[0]
    section.page_width = config["page_width"]
    section.page_height = config["page_height"]
    section.top_margin = config["margins"]["top"]
    section.bottom_margin = config["margins"]["bottom"]
    section.left_margin = config["margins"]["left"]
    section.right_margin = config["margins"]["right"]

    for style_name, size, heading in (
        ("Normal", config["body_size"], False),
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
    ):
        style = document.styles[style_name]
        style.font.name = config["heading_font_ascii"] if heading else config["body_font_ascii"]
        style.font.size = Pt(size)
        set_style_fonts(
            style,
            config["heading_font_ascii"] if heading else config["body_font_ascii"],
            config["heading_font_east_asia"] if heading else config["body_font_east_asia"],
        )

    if classification:
        paragraph = section.header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text_run(paragraph, classification, config, size_delta=-2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(footer, "Page ", config, size_delta=-2)
    add_page_field(footer.add_run())


def set_style_fonts(style: Any, ascii_font: str, east_asia_font: str) -> None:
    if style._element.rPr is None or style._element.rPr.rFonts is None:
        return
    style._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def add_page_field(run: Any) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def parse_markdown(content: str) -> list[MarkdownBlock]:
    lines = content.splitlines()
    blocks: list[MarkdownBlock] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(MarkdownBlock("paragraph", text=" ".join(line.strip() for line in paragraph_lines)))
            paragraph_lines.clear()

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            blocks.append(MarkdownBlock("heading", text=heading.group(2).strip(), level=min(len(heading.group(1)), 3)))
            index += 1
            continue

        fence = re.match(r"^```(\w+)?\s*$", stripped)
        if fence:
            flush_paragraph()
            body: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                body.append(lines[index].rstrip())
                index += 1
            if index < len(lines):
                index += 1
            blocks.append(MarkdownBlock("quote", text="\n".join(body)))
            continue

        if is_table_start(lines, index):
            flush_paragraph()
            rows = [split_table_row(lines[index])]
            index += 2
            while index < len(lines) and is_table_row(lines[index]):
                rows.append(split_table_row(lines[index]))
                index += 1
            blocks.append(MarkdownBlock("table", rows=rows))
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[index]).rstrip())
                index += 1
            blocks.append(MarkdownBlock("quote", text="\n".join(quote_lines)))
            continue

        ordered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        bullet = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if ordered or bullet:
            flush_paragraph()
            pattern = r"^\s*\d+[.)]\s+(.+)$" if ordered else r"^\s*[-*+]\s+(.+)$"
            items: list[str] = []
            while index < len(lines):
                match = re.match(pattern, lines[index])
                if not match:
                    break
                items.append(match.group(1).strip())
                index += 1
            blocks.append(MarkdownBlock("ordered_list" if ordered else "bullet_list", items=items))
            continue

        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    return blocks


def is_table_start(lines: list[str], index: int) -> bool:
    return index + 1 < len(lines) and is_table_row(lines[index]) and is_table_separator(lines[index + 1])


def is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_table_separator(line: str) -> bool:
    if not is_table_row(line):
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in split_table_row(line))


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def render_block(document: Document, block: MarkdownBlock, config: dict[str, Any]) -> None:
    if block.kind == "heading":
        add_paragraph(document, block.text, config, heading_level=block.level)
    elif block.kind == "paragraph":
        add_paragraph(document, block.text, config)
    elif block.kind == "quote":
        for line in block.text.splitlines() or [""]:
            add_paragraph(document, line, config, quote=True)
    elif block.kind == "bullet_list":
        add_list(document, block.items, config, ordered=False)
    elif block.kind == "ordered_list":
        add_list(document, block.items, config, ordered=True)
    elif block.kind == "table":
        add_table(document, block.rows, config)


def add_paragraph(
    document: Document,
    text: str,
    config: dict[str, Any],
    *,
    heading_level: int | None = None,
    quote: bool = False,
) -> None:
    style = "Normal"
    if heading_level:
        style = f"Heading {min(heading_level, 3)}"
    paragraph = document.add_paragraph(style=style)
    if heading_level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif not heading_level and not quote:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if quote:
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.right_indent = Inches(0.15)

    size = config["body_size"]
    bold_default = False
    if heading_level == 1:
        size = 16
        bold_default = True
    elif heading_level == 2:
        size = 14
        bold_default = True
    elif heading_level == 3:
        size = 12
        bold_default = True
    elif quote:
        size = max(config["body_size"] - 1, 9)
    add_inline_runs(paragraph, text, config, size_pt=size, bold_default=bold_default)
    paragraph.paragraph_format.line_spacing = config["line_spacing"]
    paragraph.paragraph_format.space_after = Pt(6)


def add_list(document: Document, items: list[str], config: dict[str, Any], *, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        paragraph = document.add_paragraph(style=style)
        add_inline_runs(paragraph, item, config, size_pt=config["body_size"])
        paragraph.paragraph_format.line_spacing = config["line_spacing"]
        paragraph.paragraph_format.space_after = Pt(3)


def add_table(document: Document, rows: list[list[str]], config: dict[str, Any]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            text = row[column_index] if column_index < len(row) else ""
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            add_inline_runs(
                paragraph,
                text,
                config,
                size_pt=max(config["body_size"] - 1, 9),
                bold_default=row_index == 0,
            )
            paragraph.paragraph_format.space_after = Pt(0)
    document.add_paragraph("")


def add_inline_runs(
    paragraph: Any,
    text: str,
    config: dict[str, Any],
    *,
    size_pt: int,
    bold_default: bool = False,
) -> None:
    cursor = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            add_text_run(paragraph, text[cursor:match.start()], config, size_pt=size_pt, bold=bold_default)
        token = match.group(0)
        value = token.strip("*_")
        bold = bold_default or token.startswith(("**", "__"))
        italic = token.startswith(("*", "_")) and not token.startswith(("**", "__"))
        if token.startswith("***"):
            bold = True
            italic = True
        add_text_run(paragraph, value, config, size_pt=size_pt, bold=bold, italic=italic)
        cursor = match.end()
    if cursor < len(text):
        add_text_run(paragraph, text[cursor:], config, size_pt=size_pt, bold=bold_default)


def add_text_run(
    paragraph: Any,
    text: str,
    config: dict[str, Any],
    *,
    size_pt: int | None = None,
    size_delta: int = 0,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run = paragraph.add_run(text)
    run.font.name = config["body_font_ascii"]
    run.font.size = Pt(size_pt if size_pt is not None else max(config["body_size"] + size_delta, 8))
    run.bold = bold
    run.italic = italic
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        run._element.rPr.rFonts.set(qn("w:ascii"), config["body_font_ascii"])
        run._element.rPr.rFonts.set(qn("w:hAnsi"), config["body_font_ascii"])
        run._element.rPr.rFonts.set(qn("w:eastAsia"), config["body_font_east_asia"])


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--language", choices=("ko", "en"), default="ko")
    parser.add_argument("--jurisdiction", choices=("korea", "us", "uk", "intl"), default="korea")
    parser.add_argument("--classification")
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--report", type=Path)
    args = parser.parse_args(argv)

    try:
        report = render_docx(
            args.input_md,
            args.output_docx,
            language=args.language,
            jurisdiction=args.jurisdiction,
            classification=args.classification,
            overwrite=args.overwrite,
            report_path=args.report,
        )
    except Exception as exc:
        print(f"FAIL: {exc}", file=sys.stderr)
        return 1

    print(json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
